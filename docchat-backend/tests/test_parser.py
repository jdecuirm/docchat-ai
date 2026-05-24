"""Smoke and edge-case tests for the document parser (Stage B)."""

from __future__ import annotations

from io import BytesIO
from pathlib import Path

import pytest
from docx import Document as DocxDocument
from pypdf import PdfWriter

from app.ingestion.parser import (
    EmptyDocumentError,
    ParsedDocument,
    ScannedPdfError,
    UnsupportedFormatError,
    parse_document,
)

FIXTURES = Path(__file__).parent / "fixtures"


def test_parse_txt_from_path_extracts_text() -> None:
    doc = parse_document(FIXTURES / "sample.txt")

    assert isinstance(doc, ParsedDocument)
    assert doc.filename == "sample.txt"
    assert len(doc.pages) == 1
    assert doc.pages[0].page_number == 1
    assert "Retrieval-Augmented Generation" in doc.pages[0].text
    assert doc.total_chars == len(doc.pages[0].text)
    assert doc.total_chars > 0


def test_parse_md_from_bytes() -> None:
    data = (FIXTURES / "sample.md").read_bytes()

    doc = parse_document(data, filename="sample.md")

    assert doc.filename == "sample.md"
    assert "DocChat AI" in doc.pages[0].text


def test_whitespace_is_normalized_but_paragraphs_preserved() -> None:
    text = parse_document(FIXTURES / "sample.txt").pages[0].text

    assert "  " not in text  # runs of spaces collapsed
    assert "\t" not in text  # tabs collapsed to single space
    assert "\n\n\n" not in text  # excess blank lines collapsed
    assert "\n\n" in text  # paragraph breaks preserved


def test_unsupported_format_raises() -> None:
    with pytest.raises(UnsupportedFormatError):
        parse_document(b"some bytes", filename="archive.zip")


def test_empty_file_raises() -> None:
    with pytest.raises(EmptyDocumentError):
        parse_document(b"", filename="empty.txt")


def test_bytes_without_filename_raises_value_error() -> None:
    with pytest.raises(ValueError, match="filename"):
        parse_document(b"data")


def test_parse_docx_preserves_paragraphs() -> None:
    buffer = BytesIO()
    document = DocxDocument()
    document.add_paragraph("First paragraph about DocChat.")
    document.add_paragraph("Second paragraph about RAG.")
    document.save(buffer)

    doc = parse_document(buffer.getvalue(), filename="generated.docx")

    assert doc.filename == "generated.docx"
    assert len(doc.pages) == 1
    assert doc.pages[0].page_number == 1
    assert "First paragraph about DocChat." in doc.pages[0].text
    assert "\n\n" in doc.pages[0].text  # paragraph break preserved


def test_scanned_pdf_without_text_layer_raises() -> None:
    writer = PdfWriter()
    writer.add_blank_page(width=200, height=200)
    buffer = BytesIO()
    writer.write(buffer)

    with pytest.raises(ScannedPdfError):
        parse_document(buffer.getvalue(), filename="scanned.pdf")
