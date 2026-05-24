"""Document parsing: turn uploaded files into clean, structured text.

Supports PDF (``pypdf``), DOCX (``python-docx``), and plain-text/Markdown
sources. The public entry point is :func:`parse_document`, which returns a
:class:`ParsedDocument` preserving per-page structure so downstream stages can
attach page numbers to citations.
"""

from __future__ import annotations

import re
import unicodedata
from io import BytesIO
from pathlib import Path

from docx import Document as DocxDocument
from pydantic import BaseModel, Field
from pypdf import PdfReader

__all__ = [
    "ParsedPage",
    "ParsedDocument",
    "parse_document",
    "DocumentParseError",
    "UnsupportedFormatError",
    "EmptyDocumentError",
    "ScannedPdfError",
]


# --- Errors -----------------------------------------------------------------
class DocumentParseError(Exception):
    """Base class for all document parsing failures."""


class UnsupportedFormatError(DocumentParseError):
    """Raised when the file extension is not a supported format."""


class EmptyDocumentError(DocumentParseError):
    """Raised when the file is empty or yields no extractable text."""


class ScannedPdfError(DocumentParseError):
    """Raised when a PDF has pages but no extractable text layer.

    This typically means the PDF is a scan or image-only document that would
    require OCR (not implemented) to read.
    """


# --- Models -----------------------------------------------------------------
class ParsedPage(BaseModel):
    """A single page of extracted text.

    Attributes:
        page_number: 1-indexed page number. Non-paginated formats (DOCX, TXT,
            MD) are represented as a single page with ``page_number == 1``.
        text: Cleaned text content of the page (may be empty for blank pages).
    """

    page_number: int = Field(ge=1)
    text: str


class ParsedDocument(BaseModel):
    """The result of parsing a document into clean, structured text.

    Attributes:
        filename: Base name of the source file.
        pages: Per-page extracted text, in document order.
        total_chars: Total number of characters across all pages.
    """

    filename: str
    pages: list[ParsedPage]
    total_chars: int


# --- Constants --------------------------------------------------------------
_TEXT_SUFFIXES = {".txt", ".md", ".markdown"}


# --- Text cleaning ----------------------------------------------------------
def _clean_text(text: str) -> str:
    """Normalize whitespace and strip control characters.

    Collapses runs of spaces/tabs and excess blank lines while preserving
    meaningful paragraph breaks (a single blank line between paragraphs).

    Args:
        text: Raw extracted text.

    Returns:
        The cleaned text.
    """
    # Normalize line endings first so control-char filtering keeps newlines.
    text = text.replace("\r\n", "\n").replace("\r", "\n")

    # Drop control/format characters (Unicode category starting with "C")
    # except newlines and tabs, which carry structure.
    text = "".join(
        ch for ch in text if ch in ("\n", "\t") or unicodedata.category(ch)[0] != "C"
    )

    # Collapse intra-line whitespace and trim each line.
    lines = [re.sub(r"[ \t]+", " ", line).strip() for line in text.split("\n")]
    text = "\n".join(lines)

    # Collapse 3+ consecutive newlines into a single blank line.
    text = re.sub(r"\n{3,}", "\n\n", text)

    return text.strip()


# --- Per-format extractors --------------------------------------------------
def _parse_pdf(data: bytes) -> list[ParsedPage]:
    """Extract text from a PDF, one :class:`ParsedPage` per page.

    Raises:
        DocumentParseError: If the PDF cannot be read (e.g. corrupt file).
        ScannedPdfError: If the PDF has pages but no extractable text.
    """
    try:
        reader = PdfReader(BytesIO(data))
        raw_pages = [page.extract_text() or "" for page in reader.pages]
    except Exception as exc:  # noqa: BLE001 - pypdf raises many error types
        raise DocumentParseError(f"Could not read PDF file: {exc}") from exc

    pages = [
        ParsedPage(page_number=index, text=_clean_text(raw))
        for index, raw in enumerate(raw_pages, start=1)
    ]

    if pages and all(not page.text for page in pages):
        raise ScannedPdfError(
            "The PDF contains no extractable text. It is likely a scanned or "
            "image-only document; OCR (not implemented) would be required."
        )

    return pages


def _parse_docx(data: bytes) -> list[ParsedPage]:
    """Extract text from a DOCX file as a single page.

    DOCX has no stable page model outside of rendering, so the whole document
    is returned as ``page_number == 1`` with paragraphs separated by blank
    lines.

    Raises:
        DocumentParseError: If the DOCX cannot be read.
    """
    try:
        document = DocxDocument(BytesIO(data))
        paragraphs = [paragraph.text for paragraph in document.paragraphs]
    except Exception as exc:  # noqa: BLE001 - python-docx raises broadly
        raise DocumentParseError(f"Could not read DOCX file: {exc}") from exc

    text = _clean_text("\n\n".join(paragraphs))
    return [ParsedPage(page_number=1, text=text)]


def _parse_text(data: bytes) -> list[ParsedPage]:
    """Decode and clean a plain-text or Markdown file as a single page."""
    try:
        text = data.decode("utf-8")
    except UnicodeDecodeError:
        text = data.decode("utf-8", errors="replace")
    return [ParsedPage(page_number=1, text=_clean_text(text))]


# --- Input loading ----------------------------------------------------------
def _load_bytes(source: str | Path | bytes, filename: str | None) -> tuple[bytes, str]:
    """Resolve ``source`` into raw bytes and a base filename.

    Args:
        source: A filesystem path (``str``/``Path``) or raw file ``bytes``.
        filename: Required when ``source`` is bytes (used for type detection);
            optional override for the resolved name when ``source`` is a path.

    Returns:
        A tuple of ``(data, base_filename)``.

    Raises:
        ValueError: If ``source`` is bytes and no ``filename`` is given.
    """
    if isinstance(source, bytes):
        if not filename:
            raise ValueError("`filename` is required when `source` is bytes.")
        return source, Path(filename).name

    path = Path(source)
    data = path.read_bytes()
    resolved = Path(filename).name if filename else path.name
    return data, resolved


# --- Public API -------------------------------------------------------------
def parse_document(
    source: str | Path | bytes,
    *,
    filename: str | None = None,
) -> ParsedDocument:
    """Parse a document into clean, structured text.

    The format is detected from the file extension. Supported types are
    ``.pdf``, ``.docx``, ``.txt``, ``.md``/``.markdown``.

    Args:
        source: A filesystem path or raw file bytes.
        filename: Required when ``source`` is bytes (drives type detection);
            otherwise an optional override for the reported filename.

    Returns:
        A :class:`ParsedDocument` with per-page text and a total character
        count.

    Raises:
        ValueError: If ``source`` is bytes without a ``filename``.
        UnsupportedFormatError: If the extension is not supported.
        EmptyDocumentError: If the file is empty or yields no text.
        ScannedPdfError: If a PDF has no extractable text layer.
        DocumentParseError: If the file is corrupt or otherwise unreadable.
    """
    data, resolved_name = _load_bytes(source, filename)
    if not data:
        raise EmptyDocumentError(f"File '{resolved_name}' is empty.")

    suffix = Path(resolved_name).suffix.lower()
    if suffix == ".pdf":
        pages = _parse_pdf(data)
    elif suffix == ".docx":
        pages = _parse_docx(data)
    elif suffix in _TEXT_SUFFIXES:
        pages = _parse_text(data)
    else:
        raise UnsupportedFormatError(
            f"Unsupported file type '{suffix or '(none)'}'. "
            "Supported formats: .pdf, .docx, .txt, .md."
        )

    total_chars = sum(len(page.text) for page in pages)
    if total_chars == 0:
        raise EmptyDocumentError(f"No text could be extracted from '{resolved_name}'.")

    return ParsedDocument(
        filename=resolved_name,
        pages=pages,
        total_chars=total_chars,
    )
